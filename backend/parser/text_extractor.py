from __future__ import annotations
import io
import re

import fitz  # PyMuPDF
from docx import Document


def extract_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_from_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    parts: list[str] = []
    for page in doc:
        parts.append(page.get_text())
    return "\n".join(parts)


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "docx":
        return extract_from_docx(file_bytes)
    elif ext == "pdf":
        return extract_from_pdf(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


# Start of a new instruction/section (don't merge these onto previous line)
_NEW_LINE_START = re.compile(
    r"^(?:Row|Rnd|Round|Next\s+(?:row|rnd|round)|CO|Cast\s+on|Sizes?|Gauge|Materials?|"
    r"Finished\s+measurements?|Abbreviations?|Notes?|\d+\.\s|[#=])",
    re.IGNORECASE,
)


def _merge_continuation_lines(text: str) -> str:
    """Merge lines that are clearly continuations of a sentence (e.g. PDF line wrap)."""
    lines = text.split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        current = lines[i].strip()
        if not current:
            i += 1
            continue
        # Append following lines that look like continuations (no sentence end, next not a new block)
        while i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if not next_line:
                i += 1
                continue
            ends_sentence = bool(re.search(r"[.?!:]\s*$", current))
            next_is_new_block = bool(_NEW_LINE_START.match(next_line))
            if ends_sentence or next_is_new_block:
                break
            current = f"{current} {next_line}"
            i += 1
        out.append(current)
        i += 1
    return "\n".join(out)


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = _merge_continuation_lines(text)
    return text.strip()
